"""
DevAssist AI — Background Celery Tasks

Clones, parses, chunks, embeds, and uploads repository content
to the vector database asynchronously.
"""

import asyncio
import shutil
import uuid
from collections.abc import Awaitable, Callable
from pathlib import Path
from typing import TypeVar

from celery.utils.log import get_task_logger

from app.models.repository import Repository
from app.providers.factory import get_embedding_provider
from app.storage.database import AsyncSessionLocal
from app.storage.vector_store import vector_store
from app.utils.clone import clone_repository
from app.utils.parser import parse_repository
from app.workers.celery_app import celery_app

logger = get_task_logger(__name__)

# Single persistent event loop per worker process. asyncio.run() would create a
# fresh loop per task and close it, breaking the module-level async_engine and
# AsyncQdrantClient singletons that stay bound to the first loop they touch.
# Reusing one loop for the process lifetime keeps those singletons valid across
# sequential prefork tasks.
_worker_loop: asyncio.AbstractEventLoop | None = None


def _get_worker_loop() -> asyncio.AbstractEventLoop:
    global _worker_loop
    if _worker_loop is None or _worker_loop.is_closed():
        _worker_loop = asyncio.new_event_loop()
        asyncio.set_event_loop(_worker_loop)
    return _worker_loop


_T = TypeVar("_T")


def _run_async(coro_factory: Callable[[], Awaitable[_T]]) -> _T:
    """Run an async coroutine on the worker's persistent event loop."""
    return _get_worker_loop().run_until_complete(coro_factory())


async def _async_index_repository(repo_id: uuid.UUID, user_id: uuid.UUID, embedding_api_key: str | None = None) -> bool:
    """Async implementation of the repository indexing pipeline."""
    async with AsyncSessionLocal() as db:
        # 1. Fetch Repository metadata
        repo = await db.get(Repository, repo_id)
        if not repo:
            logger.error(f"Repository {repo_id} not found in database.")
            return False

        if not repo.url:
            logger.error(f"Repository {repo_id} has no clone URL.")
            return False

        logger.info(f"Starting indexing for repo: {repo.name} ({repo_id})")

        try:
            # 2. Update status to Cloning
            repo.status = "cloning"
            repo.progress = 10
            await db.commit()

            # Clone repository files
            cloned_dir = await clone_repository(repo.url, str(repo_id))

            # 3. Update status to Parsing
            repo.status = "parsing"
            repo.progress = 40
            await db.commit()

            # Parse files and generate chunks
            chunks = parse_repository(cloned_dir)
            if not chunks:
                # Empty repository or no supported files
                repo.status = "completed"
                repo.progress = 100
                repo.file_count = 0
                repo.chunk_count = 0
                await db.commit()
                # Clean up empty directory
                if cloned_dir.exists():
                    shutil.rmtree(cloned_dir)
                return True

            # Deduplicate file names to get file_count
            unique_files = {chunk["file_path"] for chunk in chunks}

            # 4. Update status to Embedding
            repo.status = "embedding"
            repo.progress = 60
            await db.commit()

            # Extract text snippets to embed
            texts = [c["content"] for c in chunks]

            # Warm up and call embedding provider
            embed_provider = get_embedding_provider(api_key=embedding_api_key)

            # Embed all chunks sequentially (handles internally in Ollama)
            logger.info(f"Generating embeddings for {len(chunks)} chunks of {repo.name}...")
            embeddings = await embed_provider.embed_batch(texts)

            # 5. Update status to Uploading
            repo.status = "indexing"
            repo.progress = 85
            await db.commit()

            # Upsert into Qdrant Vector database
            logger.info("Upserting vectors into Qdrant...")
            await vector_store.upsert_code_chunks(
                user_id=user_id,
                repository_id=repo_id,
                chunks=chunks,
                embeddings=embeddings,
            )

            # 6. Mark completed
            repo.status = "completed"
            repo.progress = 100
            repo.file_count = len(unique_files)
            repo.chunk_count = len(chunks)
            await db.commit()

            # Clean up cloned files to save disk storage space
            if cloned_dir.exists():
                shutil.rmtree(cloned_dir)

            logger.info(f"Successfully completed indexing repo {repo.name}")
            return True

        except Exception as e:
            logger.error(f"Failed indexing repo {repo.name}: {str(e)}")
            # Rollback transaction and mark failed
            await db.rollback()
            repo.status = "failed"
            repo.progress = 0
            repo.error_message = str(e)
            await db.commit()

            # Clean up disk files on failure
            storage_dir = Path(cloned_dir) if "cloned_dir" in locals() else None
            if storage_dir and storage_dir.exists():
                shutil.rmtree(storage_dir)

            raise e


@celery_app.task(name="app.workers.tasks.index_repository_task", bind=True, max_retries=3)
def index_repository_task(self, repo_id_str: str, user_id_str: str, embedding_api_key: str | None = None) -> bool:
    """
    Celery background worker entry point.
    Runs the async pipeline inside an asyncio event loop.
    """
    repo_id = uuid.UUID(repo_id_str)
    user_id = uuid.UUID(user_id_str)

    # Run async pipeline on the worker's persistent event loop
    return _run_async(lambda: _async_index_repository(repo_id, user_id, embedding_api_key=embedding_api_key))


async def _async_process_document(doc_id: uuid.UUID, embedding_api_key: str | None = None) -> bool:
    """Async implementation of the document parsing and embedding pipeline."""
    from app.models.document import Document
    from app.utils.parser import chunk_file_content

    async with AsyncSessionLocal() as db:
        doc = await db.get(Document, doc_id)
        if not doc:
            logger.error(f"Document {doc_id} not found in database.")
            return False

        logger.info(f"Starting processing for document: {doc.filename} ({doc_id})")

        try:
            # 1. Read document text content
            file_path = Path(doc.storage_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Document file not found at {doc.storage_path}")

            if doc.file_type == "pdf":
                from pypdf import PdfReader

                reader = PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                content = "\n".join(text_parts)
            elif doc.file_type == "docx":
                import docx

                docx_doc = docx.Document(file_path)
                text_parts = []
                for paragraph in docx_doc.paragraphs:
                    if paragraph.text:
                        text_parts.append(paragraph.text)
                for table in docx_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)
                content = "\n".join(text_parts)
            else:
                with open(file_path, encoding="utf-8", errors="replace") as f:
                    content = f.read()

            # 2. Chunk content
            language = doc.file_type  # e.g. "md", "txt", "html", "pdf", "docx"
            chunks = chunk_file_content(doc.filename, content, language)

            if not chunks:
                doc.status = "ready"
                doc.chunk_count = 0
                await db.commit()
                return True

            # 3. Generate embeddings
            texts = [c["content"] for c in chunks]
            embed_provider = get_embedding_provider(api_key=embedding_api_key)
            embeddings = await embed_provider.embed_batch(texts)

            # 4. Upsert into Qdrant Vector database
            # Scoped to repository_id (repo_id) if available, so they are queried alongside source code
            await vector_store.upsert_code_chunks(
                user_id=doc.user_id,
                repository_id=doc.repo_id,
                chunks=chunks,
                embeddings=embeddings,
            )

            # 5. Mark ready
            doc.status = "ready"
            doc.chunk_count = len(chunks)
            await db.commit()

            logger.info(f"Successfully processed document {doc.filename}")
            return True

        except Exception as e:
            logger.error(f"Failed processing document {doc.filename}: {str(e)}")
            await db.rollback()
            doc.status = "error"
            doc.error_message = str(e)
            await db.commit()
            raise e


@celery_app.task(name="app.workers.tasks.process_document_task", bind=True, max_retries=3)
def process_document_task(self, doc_id_str: str, embedding_api_key: str | None = None) -> bool:
    """
    Celery background worker entry point for processing documents.
    """
    doc_id = uuid.UUID(doc_id_str)
    return _run_async(lambda: _async_process_document(doc_id, embedding_api_key=embedding_api_key))
